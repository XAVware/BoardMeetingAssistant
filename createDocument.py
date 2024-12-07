import json
import os
import sys
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

json_data = sys.argv[1]
month = sys.argv[2]

data = json.loads(json_data)
# print(data)

doc = Document()
doc.add_heading("Transcription", level=1)

# Set document margins (in inches)
# for section in doc.sections:
#     section.top_margin = Inches(0.5)    
#     section.bottom_margin = Inches(0.5)  
#     section.left_margin = Inches(0.5)    
#     section.right_margin = Inches(0.5)

# Set the global font by modifying the "Normal" style
style = doc.styles['Normal']
font = style.font
font.name = 'Helvetica'
font.size = Pt(11)

for utterance in data['transcript'].get("utterances", []):
    speaker = utterance.get("speakerName", "Unknown Speaker")
    text = utterance.get("text", "")

    # Create a new paragraph for each utterance
    paragraph = doc.add_paragraph()
    
    # Add speaker name in bold
    speaker_run = paragraph.add_run(f"{speaker}: ")
    speaker_run.bold = True

    # Add the text spoken by the speaker
    text_run = paragraph.add_run(text)

output_path = os.path.join(fr"/Users/ryan/Library/Containers/com.xavware.BoardMeetingAssistant/Data/Documents/{month}", 'Draft.docx')
doc.save(output_path)

print("Saved. Opening...")

path = f'open -a "/Applications/Microsoft Word.app" {output_path}'
os.system(path)